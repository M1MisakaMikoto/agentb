#!/usr/bin/env python3
"""
生成包含表格的测试 Word 文档

用于 E2E 测试表格内容提取功能
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches


def create_table_test_document(output_path: str) -> bool:
    """
    创建包含多种表格的测试文档

    Args:
        output_path: 输出文件路径

    Returns:
        bool: 是否创建成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[ERROR] 需要安装 python-docx: pip install python-docx")
        return False

    print(f"[生成] 创建表格测试文档: {output_path}")

    doc = Document()

    # 标题
    title = doc.add_heading('桥梁检测数据汇总表', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 段落说明
    intro = doc.add_paragraph(
        '本文档用于测试文档读取工具的表格提取功能。'
        '以下包含多个表格，请验证工具能否正确提取表格数据和结构。'
    )

    # 表格1：基本信息表（简单表格）
    doc.add_heading('一、桥梁基本信息', level=2)
    table1 = doc.add_table(rows=5, cols=4, style='Table Grid')
    table1_data = [
        ['项目', '数值', '单位', '备注'],
        ['桥梁名称', '朝阳寺立交桥', '-', '主桥'],
        ['全长', '1250.5', '米', '含引桥'],
        ['宽度', '24.0', '米', '双向四车道'],
        ['建成年份', '2008', '年', '-'],
    ]
    for i, row_data in enumerate(table1_data):
        for j, cell_text in enumerate(row_data):
            table1.rows[i].cells[j].text = cell_text

    # 表格2：病害统计表（多行数据）
    doc.add_heading('二、病害统计明细', level=2)
    table2 = doc.add_table(rows=7, cols=6, style='Table Grid')
    table2_data = [
        ['编号', '位置', '病害类型', '严重程度', '尺寸(cm)', '处置建议'],
        ['D001', '第3跨梁体', '纵向裂缝', 'B级', '长45×宽0.2', '封闭注浆'],
        ['D002', '5号墩顶', '混凝土剥落', 'C级', '面积30×25', '修补砂浆'],
        ['D003', '第8跨横隔板', '钢筋锈蚀', 'C级', '露筋长度15', '除锈防腐'],
        ['D004', '桥面铺装', '横向裂缝', 'A级', '长120×宽0.1', '灌缝处理'],
        ['D005', '12号支座', '变形超标', 'D级', '偏移量8', '更换支座'],
        ['D006', '锥坡护坡', '冲刷掏空', 'B级', '面积200×150', '回填压实'],
    ]
    for i, row_data in enumerate(table2_data):
        for j, cell_text in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_text

    # 表格3：BCI评分表（数值型数据）
    doc.add_heading('三、BCI评分记录', level=2)
    table3 = doc.add_table(rows=6, cols=5, style='Table Grid')
    table3_data = [
        ['检测年份', '上部结构', '下部结构', '桥面系', '综合BCI'],
        ['2019', '85.2', '88.5', '82.3', '85.3'],
        ['2020', '83.1', '87.2', '80.5', '83.6'],
        ['2021', '81.5', '86.0', '78.9', '82.1'],
        ['2022', '79.8', '84.3', '76.2', '80.1'],
        ['2023', '78.2', '82.5', '74.8', '78.5'],
    ]
    for i, row_data in enumerate(table3_data):
        for j, cell_text in enumerate(row_data):
            table3.rows[i].cells[j].text = cell_text

    # 表格4：维修建议表（带合并单元格效果）
    doc.add_heading('四、维修建议清单', level=2)
    table4 = doc.add_table(rows=5, cols=4, style='Table Grid')
    table4_data = [
        ['优先级', '工作内容', '预估费用(万元)', '计划工期'],
        ['紧急', '支座更换(D005)', '45.0', '2024Q1'],
        ['重要', '裂缝修复(D001,D004)', '18.5', '2024Q2'],
        ['一般', '混凝土修补(D002)', '8.2', '2024Q3'],
        ['观察', '定期监测(D003,D006)', '2.0', '持续'],
    ]
    for i, row_data in enumerate(table4_data):
        for j, cell_text in enumerate(row_data):
            table4.rows[i].cells[j].text = cell_text

    # 总结段落
    summary = doc.add_paragraph()
    summary.add_run('\n总结：').bold = True
    summary.add_run(
        '本报告共发现 6 处病害，其中紧急 1 处、重要 2 处、一般 2 处、观察 1 处。'
        '综合 BCI 评分为 78.5 分，属于 C 级桥梁，建议加强监测并按计划实施维修。'
    )

    # 保存文档
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    print(f'[成功] 文档已创建: {output_path} ({file_size:,} bytes)')
    print(f'       包含 4 个表格:')
    print(f'       - 表格1: 基本信息表 (5行×4列)')
    print(f'       - 表格2: 病害统计明细 (7行×6列)')
    print(f'       - 表格3: BCI评分记录 (6行×5列)')
    print(f'       - 表格4: 维修建议清单 (5行×4列)')

    return True


if __name__ == '__main__':
    # 默认输出到统一的本地 E2E fixture 目录。
    project_root = Path(__file__).resolve().parents[3]
    output_file = project_root / '.dev' / 'fixture' / 'table_test_document.docx'

    success = create_table_test_document(str(output_file))
    exit(0 if success else 1)
