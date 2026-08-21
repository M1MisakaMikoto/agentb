import os
import sys
import json
import re
import time
import tempfile
import shutil
import subprocess
import gc
from typing import Optional, Dict, Any, Tuple, List, Generator

from .registry import ToolDefinition, ToolRegistry
from .pandoc_cache import pandoc_conversation_cache
from singleton import get_settings_service

DOCUMENT_CHUNK_SIZE = 500
DOCUMENT_MAX_MEMORY_MB = 300
FORCE_GC_AFTER_CHUNKS = 3

# ============================================================
# Pandoc 转换相关
# ============================================================

def _find_pandoc() -> Optional[str]:
    """查找 pandoc 可执行文件路径"""
    import shutil

    # 检查环境变量
    for env_name in ["PANDOC_PATH", "PANDOC_HOME"]:
        env_path = os.environ.get(env_name)
        if env_path:
            candidates = [
                os.path.join(env_path, "pandoc.exe"),
                os.path.join(env_path, "pandoc"),
            ]
            for candidate in candidates:
                if os.path.isfile(candidate):
                    return candidate

    # 检查常见安装位置
    common_paths = [
        r"C:\Program Files\Pandoc\pandoc.exe",
        r"C:\Program Files (x86)\Pandoc\pandoc.exe",
    ]
    for path in common_paths:
        if os.path.isfile(path):
            return path

    # 通过 shutil.which 查找
    pandoc_path = shutil.which("pandoc")
    if pandoc_path:
        return pandoc_path

    return None


def _convert_via_pandoc(content: str, output_path: str) -> bool:
    """
    使用 pandoc 将 Markdown 内容转换为 DOCX

    Args:
        content: Markdown 格式的文本内容
        output_path: 输出的 docx 文件路径

    Returns:
        成功返回 True，失败返回 False
    """
    pandoc_path = _find_pandoc()
    if not pandoc_path:
        print("[PANDOC] 未找到 pandoc")
        return False

    try:
        import subprocess

        # 创建临时 Markdown 文件
        with tempfile.NamedTemporaryFile(
            mode='w',
            suffix='.md',
            encoding='utf-8',
            delete=False
        ) as tmp_md:
            tmp_md.write(content)
            tmp_md_path = tmp_md.name

        try:
            # 调用 pandoc 转换
            cmd = [
                pandoc_path,
                "--standalone",           # 生成完整的文档（包含页眉页脚等）
                "--from", "markdown",
                "--to", "docx",
                "-o", output_path,
                tmp_md_path
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
            )

            if result.returncode == 0:
                # 验证输出文件
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    print(f"[PANDOC] 转换成功: {output_path}")
                    return True
                else:
                    print(f"[PANDOC] 转换失败：输出文件为空或不存在")
                    return False
            else:
                print(f"[PANDOC] 转换失败: {result.stderr}")
                return False

        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_md_path)
            except Exception:
                pass

    except subprocess.TimeoutExpired:
        print("[PANDOC] 转换超时（120秒）")
        return False
    except Exception as e:
        print(f"[PANDOC] 转换异常: {e}")
        return False


def _get_ext(file_path: str) -> str:
    return os.path.splitext(file_path)[1].lower()


def _make_result(data: Optional[dict] = None, error: Optional[str] = None) -> dict:
    return {"result": data, "error": error}


def _split_content_into_chunks(content: str, chunk_size: int = DOCUMENT_CHUNK_SIZE) -> Generator[List[str], None, None]:
    """将内容按行分割成块，降低峰值内存"""
    lines = content.split('\n')
    current_chunk = []
    
    for line in lines:
        if line.strip():
            current_chunk.append(line)
        
        if len(current_chunk) >= chunk_size:
            yield current_chunk
            current_chunk = []
    
    if current_chunk:
        yield current_chunk


def _check_memory_before_operation(threshold_mb: int = DOCUMENT_MAX_MEMORY_MB) -> bool:
    """执行操作前检查内存，超阈值主动释放"""
    try:
        import psutil
        process = psutil.Process(os.getpid())
        memory_mb = process.memory_info().rss / 1024 / 1024
        
        if memory_mb > threshold_mb:
            print(f"[MEMORY] 操作前内存偏高: {memory_mb:.1f}MB > {threshold_mb}MB")
            collected = gc.collect()
            
            memory_after = process.memory_info().rss / 1024 / 1024
            print(f"[MEMORY] GC 后: {memory_after:.1f}MB (释放: {memory_mb - memory_after:.1f}MB)")
            
            return memory_after <= threshold_mb * 1.5
        
        return True
    except ImportError:
        return True


# ============================================================
# PDF 操作 (r/w/a/u)
# ============================================================

def _pdf_read(file_path: str, start_idx: int = 0, max_length: int = 100000,
              include_metadata: bool = True, use_llm_parsing: bool = True) -> dict:
    try:
        import pypdf
    except ImportError:
        return _make_result(error="缺少依赖: pip install pypdf")
    
    try:
        if use_llm_parsing:
            try:
                import pymupdf4llm
            except ImportError:
                return _make_result(error="缺少依赖: pip install pymupdf4llm")
            md_text = pymupdf4llm.to_markdown(file_path)
            full_text = md_text
        else:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                full_text = "\n".join(text_parts)

        total_length = len(full_text)

        end_idx = min(start_idx + max_length, total_length)
        content = full_text[start_idx:end_idx]

        metadata = {}
        if include_metadata:
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    metadata = {
                        "file_type": "pdf",
                        "page_count": len(reader.pages),
                        "author": reader.metadata.author if reader.metadata else None,
                        "title": reader.metadata.title if reader.metadata else None,
                        "creator": reader.metadata.creator if reader.metadata else None,
                        "parsing_mode": "llm" if use_llm_parsing else "fast",
                    }
            except Exception:
                metadata = {"file_type": "pdf", "page_count": "unknown"}
        
        return _make_result({
            "content": content,
            "metadata": metadata,
            "total_length": total_length,
            "read_range": f"{start_idx}-{end_idx}",
            "truncated": end_idx < total_length
        })
    except Exception as e:
        return _make_result(error=f"PDF读取失败: {str(e)}")


def _pdf_write(file_path: str, content: str, metadata: Optional[dict] = None) -> dict:
    """将 Markdown 直接渲染为 PDF（WeasyPrint：HTML/CSS 排版，支持中文）。"""
    try:
        from .pdf_renderer import render_markdown_to_pdf
        result = render_markdown_to_pdf(content, file_path, metadata)
        return _make_result(result)
    except Exception as e:
        return _make_result(error=f"PDF写入失败: {str(e)}")


def _pdf_append(file_path: str, content: str) -> dict:
    """PDF 不支持追加，明确报错（不写兜底）

    PDF 格式不适合增量追加，建议用 document w 重新生成完整 PDF。
    """
    return _make_result(
        error="PDF暂不支持追加操作，请用 document w 重新生成完整PDF"
    )


def _pdf_update(file_path: str, target: str, content: Optional[str] = None,
                field: Optional[str] = None) -> dict:
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return _make_result(error="缺少依赖: pip install pypdf")
    
    try:
        if field == "metadata":
            reader = PdfReader(file_path)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            
            meta_data = {}
            if isinstance(content, str):
                try:
                    meta_data = json.loads(content)
                except json.JSONDecodeError:
                    meta_data = {"info_string": content}
            elif isinstance(content, dict):
                meta_data = content
            
            if meta_data:
                from pypdf.generic import NameObject, TextStringObject, create_string_object
                info_dict = {}
                for k, v in meta_data.items():
                    info_dict[NameObject(f"/{k}")] = create_string_object(str(v))
                if info_dict:
                    writer.add_metadata(info_dict)
            
            with open(file_path, "wb") as f:
                writer.write(f)
            return _make_result({"message": "PDF元数据更新成功"})
        
        return _make_result(error=f"不支持的更新操作: field={field}, 支持修改元数据(metadata)")
    except Exception as e:
        return _make_result(error=f"PDF更新失败: {str(e)}")


# ============================================================
# DOCX / DOC 操作 (r/w/a/u)
# ============================================================

def _convert_doc_to_docx(file_path: str) -> Optional[str]:
    """Convert .doc to .docx using multiple methods with retry logic."""
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            temp_docx = tempfile.mktemp(suffix=".docx")
            
            # Method 1: Win32 COM (Windows only, most reliable)
            if sys.platform == "win32":
                try:
                    import win32com.client
                    import pythoncom
                    
                    pythoncom.CoInitialize()
                    
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    doc = word.Documents.Open(
                        os.path.abspath(file_path),
                        Visible=False,
                        ConfirmConversions=False
                    )
                    
                    # Save as docx (wdFormatXMLDocument = 16)
                    doc.SaveAs2(os.path.abspath(temp_docx), FileFormat=16)
                    doc.Close()
                    word.Quit()
                    
                    pythoncom.CoUninitialize()
                    
                    if os.path.exists(temp_docx):
                        return temp_docx
                        
                except ImportError:
                    pass
                except Exception as com_error:
                    if attempt < max_retries - 1:
                        time.sleep(2 * (attempt + 1))
                        continue
            
            # Method 2: docx2python library
            try:
                from docx2python import docx2python
                docx2python(file_path, temp_docx)
                if os.path.exists(temp_docx):
                    return temp_docx
            except ImportError:
                pass
            
            # Method 3: LibreOffice command line
            try:
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx", "--outdir",
                     os.path.dirname(temp_docx), file_path],
                    capture_output=True, timeout=30
                )
                if result.returncode == 0:
                    output_dir = os.path.dirname(temp_docx)
                    output_name = os.path.basename(file_path).rsplit(".", 1)[0] + ".docx"
                    converted = os.path.join(output_dir, output_name)
                    if os.path.exists(converted):
                        shutil.move(converted, temp_docx)
                        return temp_docx
            except FileNotFoundError:
                pass
            
            break
            
        except Exception:
            if attempt < max_retries - 1:
                time.sleep(2 * (attempt + 1))
                continue
    
    return None


def _extract_table_data(table) -> Tuple[str, dict]:
    """
    提取单个表格的数据，返回 Markdown 格式文本和元数据

    Args:
        table: python-docx 的 Table 对象

    Returns:
        Tuple[str, dict]: (Markdown 格式表格字符串, 元数据字典)
    """
    rows_data = []
    max_cols = 0

    # 遍历表格行
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            row_cells.append(cell_text)

        # 记录最大列数（处理合并单元格）
        if len(row_cells) > max_cols:
            max_cols = len(row_cells)
        rows_data.append(row_cells)

    if not rows_data or max_cols == 0:
        return "", {"rows": 0, "cols": 0}

    # 格式化为 Markdown 表格
    lines = []

    # 表头行（第一行）
    header = rows_data[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")

    # 数据行
    for row in rows_data[1:]:
        # 补齐列数（处理某些行列数不一致的情况）
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")

    formatted_text = "\n".join(lines)

    metadata = {
        "rows": len(rows_data),
        "cols": max_cols,
    }

    return formatted_text, metadata


def _build_docx_read_result(
    file_path: str,
    full_text: str,
    start_idx: int,
    max_length: int,
    include_metadata: bool,
    method: str,
    structure: Optional[list] = None,
    metadata_extra: Optional[dict] = None,
    offset: Optional[int] = None,
    limit: Optional[int] = None,
) -> dict:
    total_length = len(full_text)
    total_lines = len(full_text.splitlines())
    file_size = os.path.getsize(file_path)
    line_mode = offset is not None or limit is not None
    if line_mode:
        assert offset is not None and limit is not None
        assert offset >= 1
        assert limit >= 1
        raw_lines = full_text.splitlines(keepends=True)
        if offset > total_lines:
            start_idx = total_length
            end_idx = total_length
            content = ""
            start_line = total_lines + 1
            end_line = total_lines
        else:
            start_line = offset
            end_line = min(offset + limit - 1, total_lines)
            start_idx = sum(len(line) for line in raw_lines[: offset - 1])
            page_lines = raw_lines[offset - 1 : end_line]
            content = "".join(page_lines)
            end_idx = start_idx + len(content)
    else:
        if start_idx >= total_length and total_length > 0:
            end_idx = total_length
            content = ""
            start_line = total_lines + 1
            end_line = total_lines
        else:
            requested_end = min(start_idx + max_length, total_length)
            end_idx = requested_end
            if requested_end < total_length:
                last_newline = full_text.rfind("\n", start_idx, requested_end)
                if last_newline >= start_idx:
                    end_idx = last_newline + 1
            content = full_text[start_idx:end_idx]
            start_line = full_text.count("\n", 0, start_idx) + 1
            returned_line_count = len(content.splitlines())
            end_line = (
                start_line + returned_line_count - 1
                if returned_line_count
                else start_line - 1
            )

    metadata = {}
    if include_metadata:
        metadata = {
            "file_type": str(_get_ext(file_path).lstrip(".")),
            "method": method,
            **(metadata_extra or {}),
        }

    if (line_mode and offset > total_lines) or (
        not line_mode and start_idx >= total_length and total_length > 0
    ):
        result = {
            "content": "",
            "metadata": metadata,
            "structure": structure or [],
            "total_length": total_length,
            "total_lines": total_lines,
            "start_line": total_lines + 1,
            "end_line": total_lines,
            "file_size": file_size,
            "read_range": f"{total_length}-{total_length}",
            "pagination_mode": "lines" if line_mode else "characters",
            "truncated": False,
            "status": "end_of_file",
            "message": f"Document fully read (total length: {total_length})",
        }
        if line_mode:
            result["next_offset"] = None
        else:
            result["next_start_idx"] = None
        return result

    result = {
        "content": content,
        "metadata": metadata,
        "structure": structure or [],
        "total_length": total_length,
        "total_lines": total_lines,
        "start_line": start_line,
        "end_line": end_line,
        "file_size": file_size,
        "read_range": f"{start_idx}-{end_idx}",
        "pagination_mode": "lines" if line_mode else "characters",
    }
    if line_mode:
        result["truncated"] = end_line < total_lines
        result["next_offset"] = end_line + 1 if result["truncated"] else None
    else:
        result["truncated"] = end_idx < total_length
        result["next_start_idx"] = end_idx if result["truncated"] else None
    return result


def _build_pandoc_read_result(
    file_path: str,
    full_text: str,
    start_idx: int,
    max_length: int,
    include_metadata: bool,
    offset: Optional[int] = None,
    limit: Optional[int] = None,
) -> dict:
    return _build_docx_read_result(
        file_path,
        full_text,
        start_idx,
        max_length,
        include_metadata,
        method="pandoc",
        offset=offset,
        limit=limit,
    )


def _docx_read_via_pandoc(
    file_path: str,
    start_idx: int = 0,
    max_length: int = 100000,
    include_metadata: bool = True,
    conversation_id: Optional[str] = None,
    offset: Optional[int] = None,
    limit: Optional[int] = None,
) -> Optional[dict]:
    """
    使用 Pandoc 读取 Word 文档（高覆盖率方案）

    优势：
    - 自动提取：段落、表格、图片、页眉页脚、脚注、超链接等
    - 输出格式：Markdown，保留文档结构
    - 覆盖率：95%+ 的文档元素

    Args:
        file_path: 文件路径
        start_idx: 起始位置（字符偏移）
        max_length: 最大读取长度
        include_metadata: 是否包含元数据

    Returns:
        成功返回 dict，失败返回 None（调用方应降级到 python-docx）
    """
    if conversation_id:
        def load_uncached() -> Optional[str]:
            uncached = _docx_read_via_pandoc(
                file_path,
                start_idx=0,
                max_length=sys.maxsize,
                include_metadata=False,
            )
            return uncached["content"] if uncached is not None else None

        full_text = pandoc_conversation_cache.get_or_load(
            conversation_id,
            file_path,
            load_uncached,
        )
        if full_text is None:
            return None
        return _build_pandoc_read_result(
            file_path,
            full_text,
            start_idx,
            max_length,
            include_metadata,
            offset,
            limit,
        )

    pandoc_path = _find_pandoc()
    if not pandoc_path:
        print("[DOCX-READ] ⚠️ Pandoc 未安装，将降级到 python-docx 模式")
        print("[DOCX-READ] ⚠️ 降级原因: 无法找到 pandoc 可执行文件")
        print("[DOCX-READ] ⚠️ 影响范围: 图片、页眉页脚、脚注等元素可能丢失")
        return None

    try:
        import subprocess

        actual_path = file_path
        cleanup = False

        # 处理 .doc 格式
        if _get_ext(file_path) == ".doc":
            converted = _convert_doc_to_docx(file_path)
            if not converted:
                print("[DOCX-READ] ⚠️ .doc 格式转换失败，尝试使用 pandoc 直接读取")
                # pandoc 可以直接读取 .doc，不需要转换
            else:
                actual_path = converted
                cleanup = True

        # 使用 pandoc 转换为 Markdown
        cmd = [
            pandoc_path,
            "--from=docx",
            "--to=markdown",
            "--wrap=none",           # 不自动换行
            "--extract-media=./",    # 提取内嵌媒体（可选）
            actual_path
        ]

        print(f"[DOCX-READ] 使用 Pandoc 读取文档: {file_path}")

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=60,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        )

        if result.returncode != 0:
            print(f"[DOCX-READ] ❌ Pandoc 读取失败: {result.stderr}")
            print("[DOCX-READ] ⚠️ 将降级到 python-docx 模式")
            return None

        full_text = result.stdout

        # 清理临时文件
        if cleanup and actual_path != file_path and os.path.exists(actual_path):
            os.unlink(actual_path)

        print(f"[DOCX-READ] ✅ Pandoc 读取成功 (方法: pandoc, 长度: {len(full_text)})")
        return _build_pandoc_read_result(
            file_path,
            full_text,
            start_idx,
            max_length,
            include_metadata,
            offset,
            limit,
        )

    except subprocess.TimeoutExpired:
        print("[DOCX-READ] ❌ Pandoc 读取超时（60秒），降级到 python-docx")
        return None
    except Exception as e:
        print(f"[DOCX-READ] ❌ Pandoc 异常: {str(e)}")
        print("[DOCX-READ] ⚠️ 将降级到 python-docx 模式")
        return None


def _docx_read(
    file_path: str,
    start_idx: int = 0,
    max_length: int = 100000,
    include_metadata: bool = True,
    conversation_id: Optional[str] = None,
    offset: Optional[int] = None,
    limit: Optional[int] = None,
) -> dict:
    """
    读取 Word 文档（双模式：Pandoc 优先 + python-docx 降级）

    策略：
    1. 优先使用 Pandoc（高覆盖率，支持图片/页眉页脚/脚注等）
    2. Pandoc 不可用时，降级到 python-docx（基础覆盖，段落+表格）
    3. 所有降级操作都会在控制台和日志中留痕
    """
    # ===== 模式 1：尝试使用 Pandoc（高覆盖率）=====
    print(f"[DOCX-READ] 开始读取文档: {file_path}")
    print("[DOCX-READ] 策略: 优先使用 Pandoc (高覆盖率模式)")

    pandoc_result = _docx_read_via_pandoc(
        file_path,
        start_idx,
        max_length,
        include_metadata,
        conversation_id,
        offset,
        limit,
    )

    if pandoc_result is not None:
        # Pandoc 成功
        return _make_result(pandoc_result)

    # ===== 模式 2：降级到 python-docx =====
    print("=" * 60)
    print("[DOCX-READ] ⚠️ ⚠️ ⚠️ 降级警告 ⚠️ ⚠️ ⚠️")
    print("[DOCX-READ] 已从 Pandoc 模式降级到 python-docx 模式")
    print("[DOCX-READ] 降级影响:")
    print("   - ❌ 内嵌图片可能丢失")
    print("   - ❌ 页眉页脚内容可能丢失")
    print("   - ❌ 脚注尾注可能丢失")
    print("   - ❌ 批注注释可能丢失")
    print("   - ✅ 段落文本正常提取")
    print("   - ✅ 表格数据正常提取")
    print(f"[DOCX-READ] 文档路径: {file_path}")
    print("=" * 60)

    try:
        from docx import Document
    except ImportError:
        return _make_result(error="缺少依赖: pip install python-docx")
    
    try:
        actual_path = file_path
        cleanup = False
        
        if _get_ext(file_path) == ".doc":
            converted = _convert_doc_to_docx(file_path)
            if not converted:
                return _make_result(error=".doc格式转换失败，请安装LibreOffice或docx2python")
            actual_path = converted
            cleanup = True
        
        doc = Document(actual_path)

        # 统一存储段落和表格内容，保持文档顺序
        content_parts = []
        structure = []

        # 1. 遍历段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                content_parts.append(text)
                style_name = para.style.name if para.style else None
                structure.append({
                    "type": "paragraph",
                    "index": len(content_parts) - 1,
                    "style": str(style_name) if style_name is not None else None,
                    "content": text[:100] + "..." if len(text) > 100 else text
                })

        # 2. 遍历表格
        for t_idx, table in enumerate(doc.tables):
            table_text, table_meta = _extract_table_data(table)
            if table_text.strip():
                content_parts.append(table_text)
                structure.append({
                    "type": "table",
                    "index": t_idx,
                    "rows": table_meta["rows"],
                    "cols": table_meta["cols"],
                    "content_preview": table_text[:100] + "..." if len(table_text) > 100 else table_text
                })

        full_text = "\n\n".join(content_parts)
        total_length = len(full_text)
        paragraph_count = sum(1 for item in structure if item["type"] == "paragraph")
        table_count = sum(1 for item in structure if item["type"] == "table")


        metadata = {}
        if include_metadata:
            core_props = doc.core_properties
            # 统计段落数量和表格数量
            paragraph_count = sum(1 for item in structure if item["type"] == "paragraph")
            table_count = sum(1 for item in structure if item["type"] == "table")

            metadata = {
                "file_type": str(_get_ext(file_path).lstrip(".")),
                "method": "python-docx",  # 标记降级模式
                "paragraph_count": int(paragraph_count),
                "table_count": int(table_count),
                "author": str(core_props.author) if core_props.author else None,
                "title": str(core_props.title) if core_props.title else None,
                "subject": str(core_props.subject) if core_props.subject else None,
            }

        if cleanup and actual_path != file_path:
            os.unlink(actual_path)

        # 降级模式成功日志
        print(f"[DOCX-READ] ✅ python-docx 降级模式读取成功")
        print(f"[DOCX-READ] 📊 统计: {paragraph_count} 个段落, {table_count} 个表格, 总长度 {total_length} 字符")
        print("[DOCX-READ] 💡 提示: 安装 Pandoc 可获得更完整的文档内容（含图片/页眉页脚/脚注等）")

        return _make_result(_build_docx_read_result(
            file_path,
            full_text,
            start_idx,
            max_length,
            include_metadata,
            method="python-docx",
            structure=structure[:20],
            metadata_extra={
                key: value
                for key, value in metadata.items()
                if key not in {"file_type", "method"}
            },
            offset=offset,
            limit=limit,
        ))
    except Exception as e:
        print(f"[DOCX-READ] ❌ python-docx 降级模式也失败了: {str(e)}")
        return _make_result(error=f"Word文档读取失败: {str(e)}")


def _search_canonical_text(
    file_path: str,
    full_text: str,
    pattern: str,
    case_sensitive: bool = False,
    context: int = 2,
    max_results: int = 50,
    start_idx: int = 0,
) -> dict:
    """Search canonical text with grep-like, one-result-per-line semantics."""
    flags = 0 if case_sensitive else re.IGNORECASE
    try:
        regex = re.compile(pattern, flags)
    except re.error as exc:
        return _make_result(error=f"正则表达式错误: {exc}")

    try:
        context = max(0, int(context))
        max_results = max(1, int(max_results))
        start_idx = int(start_idx)
    except (TypeError, ValueError):
        return _make_result(error="context、max_results 和 start_idx 必须是整数")
    if start_idx < 0:
        return _make_result(error="start_idx 不能小于 0")

    segments = []
    offset = 0
    for segment_index, raw_line in enumerate(full_text.splitlines(keepends=True)):
        line = raw_line.rstrip("\r\n")
        segments.append({
            "segment_index": segment_index,
            "line": line,
            "char_start": offset,
            "char_end": offset + len(line),
            "raw_end": offset + len(raw_line),
        })
        offset += len(raw_line)

    if full_text and not segments:
        segments.append({
            "segment_index": 0,
            "line": full_text,
            "char_start": 0,
            "char_end": len(full_text),
            "raw_end": len(full_text),
        })

    matching_segments = []
    total_occurrences = 0
    for segment in segments:
        line = segment["line"]
        if not line.strip():
            continue
        occurrences = []
        segment_index = segment["segment_index"]
        for match in regex.finditer(line):
            char_start = segment["char_start"] + match.start()
            char_end = segment["char_start"] + match.end()
            total_occurrences += 1
            occurrences.append({
                "matched_text": match.group(),
                "char_start": char_start,
                "char_end": char_end,
                "segment_index": segment_index,
                "segment_number": segment_index + 1,
                "position_in_segment": match.start(),
                "position": match.start(),
            })

        if occurrences:
            matching_segments.append((segment, occurrences))

    page_candidates = [
        (segment, occurrences)
        for segment, occurrences in matching_segments
        if any(occurrence["char_start"] >= start_idx for occurrence in occurrences)
    ]
    page = page_candidates[:max_results]
    matches = []
    for segment, occurrences in page:
        visible_occurrences = [
            occurrence
            for occurrence in occurrences
            if occurrence["char_start"] >= start_idx
        ]
        first = visible_occurrences[0]
        segment_index = segment["segment_index"]
        context_start_index = max(0, segment_index - context)
        context_end_index = min(len(segments) - 1, segment_index + context)
        read_start = segments[context_start_index]["char_start"]
        read_end = segments[context_end_index]["raw_end"]
        index = {
            "char_start": first["char_start"],
            "char_end": first["char_end"],
            "segment_index": segment_index,
            "segment_number": segment_index + 1,
        }
        matches.append({
            "pattern": pattern,
            "matched_text": first["matched_text"],
            "matched_texts": [
                occurrence["matched_text"] for occurrence in visible_occurrences
            ],
            "occurrences": visible_occurrences,
            "snippet": full_text[read_start:read_end],
            "line": segment["line"],
            "index": index,
            **index,
            "segment_type": "line",
            "position_in_segment": first["position_in_segment"],
            "position": first["position"],
            "read_hint": {
                "start_idx": read_start,
                "max_length": max(1, read_end - read_start),
            },
        })

    next_start_idx = None
    if len(page_candidates) > len(page) and page:
        next_start_idx = page[-1][0]["raw_end"]

    return _make_result({
        "matches": matches,
        "total_matches": len(matching_segments),
        "total_occurrences": total_occurrences,
        "returned_matches": len(matches),
        "search_start_idx": start_idx,
        "next_start_idx": next_start_idx,
        "pattern": pattern,
        "file": file_path,
        "index_unit": "character",
        "segment_unit": "line",
        "total_length": len(full_text),
    })


def _docx_search(
    file_path: str,
    pattern: str,
    case_sensitive: bool = False,
    context: int = 2,
    max_results: int = 50,
    conversation_id: Optional[str] = None,
    start_idx: int = 0,
) -> dict:
    """Search the same canonical Word text used by document(r)."""
    read_result = _docx_read(
        file_path,
        start_idx=0,
        max_length=sys.maxsize,
        include_metadata=False,
        conversation_id=conversation_id,
    )
    if read_result.get("error") is not None:
        return _make_result(error=f"Word文档搜索失败: {read_result['error']}")
    full_text = (read_result.get("result") or {}).get("content", "")
    return _search_canonical_text(
        file_path,
        full_text,
        pattern,
        case_sensitive,
        context,
        max_results,
        start_idx,
    )


def _pdf_search(file_path: str, pattern: str, case_sensitive: bool = False,
                context: int = 2, max_results: int = 50, use_llm_parsing: bool = True,
                start_idx: int = 0) -> dict:
    """Search the same canonical PDF text used by document(r)."""
    read_result = _pdf_read(
        file_path,
        start_idx=0,
        max_length=sys.maxsize,
        include_metadata=False,
        use_llm_parsing=use_llm_parsing,
    )
    if read_result.get("error") is not None:
        return _make_result(error=f"PDF搜索失败: {read_result['error']}")
    full_text = (read_result.get("result") or {}).get("content", "")
    return _search_canonical_text(
        file_path,
        full_text,
        pattern,
        case_sensitive,
        context,
        max_results,
        start_idx,
    )


def _excel_search(file_path: str, pattern: str, case_sensitive: bool = False,
                  context: int = 2, max_results: int = 50,
                  start_idx: int = 0) -> dict:
    """Search the same canonical spreadsheet text used by document(r)."""
    read_result = _excel_read(
        file_path,
        start_idx=0,
        max_length=sys.maxsize,
        include_metadata=True,
    )
    if read_result.get("error") is not None:
        return _make_result(error=f"Excel搜索失败: {read_result['error']}")
    data = read_result.get("result") or {}
    result = _search_canonical_text(
        file_path,
        data.get("content", ""),
        pattern,
        case_sensitive,
        context,
        max_results,
        start_idx,
    )
    if result.get("error") is None:
        result["result"]["sheets_searched"] = (
            data.get("metadata") or {}
        ).get("sheet_names", [])
    return result
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.name = 'SimSun'
    _set_east_asia_font(run, 'SimSun')
    run.font.size = __import__('docx.shared').Pt(10.5)


def _set_east_asia_font(run, font_name: str):
    """安全地设置中文字体（兼容不同版本的 python-docx）"""
    try:
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass


def _check_memory_usage(threshold_mb: int = 512) -> bool:
    """检查当前进程内存使用情况，超过阈值主动释放"""
    try:
        import psutil
        import gc
        import os
        
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        memory_mb = memory_info.rss / 1024 / 1024
        
        if memory_mb > threshold_mb:
            print(f"[WARNING] 内存使用过高: {memory_mb:.1f}MB > {threshold_mb}MB")
            print("[ACTION] 执行垃圾回收...")
            
            collected = gc.collect()
            print(f"[GC] 回收了 {collected} 个对象")
            
            memory_info_after = process.memory_info()
            memory_mb_after = memory_info_after.rss / 1024 / 1024
            print(f"[MEMORY] GC 后内存: {memory_mb_after:.1f}MB (释放: {memory_mb - memory_mb_after:.1f}MB)")
            
            if memory_mb_after > threshold_mb * 1.5:
                raise MemoryError(f"内存严重不足: {memory_mb_after:.1f}MB (阈值: {threshold_mb}MB)")
            
            return False
        return True
    except ImportError:
        return True


def _set_table_borders(table):
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    except Exception:
        pass


def _add_cover_page(doc, title, subtitle=None, date=None, company=None):
    try:
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return
    
    for _ in range(6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = 'SimHei'
    _set_east_asia_font(run, 'SimHei')
    
    if subtitle:
        doc.add_paragraph()
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.name = 'SimSun'
        _set_east_asia_font(run, 'SimSun')
    
    for _ in range(8):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_lines = []
    if date:
        info_lines.append(date)
    if company:
        info_lines.append(company)
    
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = 'SimSun'
        _set_east_asia_font(run, 'SimSun')
    
    doc.add_page_break()


def _markdown_to_docx_content(markdown_text: str, doc, chunk_size: int = 500) -> None:
    """
    分块处理 Markdown 内容，避免大文档导致内存爆炸
    
    Args:
        markdown_text: 完整的 Markdown 文本
        doc: python-docx Document 对象
        chunk_size: 每个块的最大行数（默认500行）
    """
    try:
        import re
        import gc
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        pass
    
    lines = markdown_text.split("\n")
    total_lines = len(lines)
    
    print(f"[DOCX] 开始处理文档，总行数: {total_lines}")
    
    i = 0
    in_cover = False
    cover_data = {}
    processed_chunks = 0
    
    while i < len(lines):
        line = lines[i]
        
        if '{{cover' in line and '}}' in line:
            in_cover = True
            match = re.search(r'\{\{cover:(\w+)\}\}\s*(.*)', line)
            if match:
                cover_data[match.group(1)] = match.group(2).strip()
            i += 1
            continue
        
        if in_cover and line.strip() == '' and i + 1 < len(lines) and not lines[i+1].startswith('{{'):
            in_cover = False
            if cover_data.get('title'):
                _add_cover_page(doc,
                    title=cover_data.get('title', ''),
                    subtitle=cover_data.get('subtitle'),
                    date=cover_data.get('date'),
                    company=cover_data.get('company'))
            cover_data = {}
        
        if in_cover:
            match = re.search(r'\{\{cover:(\w+)\}\}\s*(.*)', line)
            if match:
                cover_data[match.group(1)] = match.group(2).strip()
            i += 1
            continue
        
        if line.startswith("# "):
            text = line[2:].strip()
            heading = doc.add_heading(level=1)
            run = heading.add_run(text)
            run.font.name = 'SimHei'
            _set_east_asia_font(run, 'SimHei')
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 10 else None
        elif line.startswith("## "):
            text = line[3:].strip()
            heading = doc.add_heading(level=2)
            run = heading.add_run(text)
            run.font.name = 'SimHei'
            _set_east_asia_font(run, 'SimHei')
        elif line.startswith("### "):
            text = line[4:].strip()
            heading = doc.add_heading(level=3)
            run = heading.add_run(text)
            run.font.name = 'SimHei'
            _set_east_asia_font(run, 'SimHei')
        elif line.startswith("#### "):
            text = line[5:].strip()
            heading = doc.add_heading(level=4)
            run = heading.add_run(text)
            run.font.name = 'SimSun'
            _set_east_asia_font(run, 'SimSun')
        elif line.startswith("---") or line.startswith("***"):
            pass
        elif line.startswith("- ") or line.startswith("* "):
            items = []
            while i < len(lines) and (lines[i].startswith("- ") or lines[i].startswith("* ")):
                items.append(lines[i][2:].strip())
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style='List Bullet')
                for run in p.runs:
                    run.font.name = 'SimSun'
                    _set_east_asia_font(run, 'SimSun')
                    run.font.size = Pt(12)
            continue
        elif re.match(r'^\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                items.append(re.sub(r'^\d+\.\s', '', lines[i]).strip())
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style='List Number')
                for run in p.runs:
                    run.font.name = 'SimSun'
                    _set_east_asia_font(run, 'SimSun')
                    run.font.size = Pt(12)
            continue
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 3:
                header_line = table_lines[0]
                headers = [cell.strip() for cell in header_line.strip("|").split("|")]
                
                rows = []
                for tl in table_lines[2:]:
                    if "---" not in tl:
                        row_cells = [cell.strip() for cell in tl.strip("|").split("|")]
                        rows.append(row_cells)
                
                if headers:
                    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                    table.style = 'Table Grid'
                    _set_table_borders(table)
                    
                    for j, h in enumerate(headers):
                        if j < len(table.rows[0].cells):
                            _set_cell_text(table.rows[0].cells[j], h, bold=True)
                    
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell in enumerate(row):
                            if r_idx + 1 < len(table.rows) and c_idx < len(table.rows[r_idx+1].cells):
                                _set_cell_text(table.rows[r_idx+1].cells[c_idx], cell)
            continue
        elif line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Cm(0.74) if any('\u4e00' <= c <= '\u9fff' for c in line) else None
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.name = 'SimSun'
                _set_east_asia_font(run, 'SimSun')
                run.font.size = Pt(12)
        
        i += 1
        
        processed_chunks += 1
        
        if processed_chunks % chunk_size == 0:
            progress_pct = (i / total_lines) * 100
            print(f"[DOCX] 处理进度: {progress_pct:.1f}% ({i}/{total_lines})")
            
            if not _check_memory_usage():
                print("[WARNING] 内存使用较高，尝试垃圾回收...")
                gc.collect()
    
    print(f"[DOCX] 文档处理完成，共处理 {i} 行")


def _docx_write(file_path: str, content: str, metadata: Optional[dict] = None) -> dict:
    """写入 Word 文档，优先使用 pandoc 转换 Markdown，然后回退到 python-docx"""

    ext = _get_ext(file_path)
    target_path = file_path if ext == ".docx" else tempfile.mktemp(suffix=".docx")

    # 尝试使用 pandoc 转换
    print(f"[DOCX] 尝试使用 pandoc 转换 Markdown...")
    pandoc_success = _convert_via_pandoc(content, target_path)

    if pandoc_success:
        # pandoc 转换成功
        if ext == ".doc":
            return _convert_docx_to_doc(target_path, file_path)

        return _make_result({"message": f"Word文档创建成功: {file_path}", "method": "pandoc"})

    # pandoc 不可用，回退到 python-docx
    print(f"[DOCX] pandoc 不可用，回退到 python-docx 模式...")

    try:
        from docx import Document
    except ImportError:
        return _make_result(error="缺少依赖: pip install python-docx")

    try:
        if not _check_memory_before_operation():
            return _make_result(error=f"内存不足，当前使用超过 {DOCUMENT_MAX_MEMORY_MB}MB")

        doc = Document()

        meta = metadata or {}
        if meta.get("author"):
            doc.core_properties.author = meta["author"]
        if meta.get("title"):
            doc.core_properties.title = meta["title"]
        if meta.get("subject"):
            doc.core_properties.subject = meta["subject"]

        total_chars = len(content)
        print(f"[DOCX-CHUNKED] 开始分块处理文档 (总字符: {total_chars:,})")

        chunks = list(_split_content_into_chunks(content))
        total_chunks = len(chunks)
        print(f"[DOCX-CHUNKED] 分为 {total_chunks} 个块 (每块 ~{DOCUMENT_CHUNK_SIZE} 行)")

        processed_chunks = 0

        for chunk_idx, chunk in enumerate(chunks):
            for line in chunk:
                if line.strip():
                    p = doc.add_paragraph(line)
                    try:
                        from docx.shared import Cm, Pt
                        p.paragraph_format.first_line_indent = Cm(0.74) if any('一' <= c <= '鿿' for c in line) else None
                        p.paragraph_format.line_spacing = 1.5
                        for run in p.runs:
                            run.font.name = 'SimSun'
                            _set_east_asia_font(run, 'SimSun')
                            run.font.size = Pt(12)
                    except Exception:
                        pass

            processed_chunks += 1

            if processed_chunks % FORCE_GC_AFTER_CHUNKS == 0:
                progress_pct = (processed_chunks / total_chunks) * 100
                print(f"[DOCX-CHUNKED] 处理进度: {progress_pct:.1f}% (块 {processed_chunks}/{total_chunks})")

                collected = gc.collect()
                if collected > 0:
                    print(f"[DOCX-CHUNKED] 执行 GC 回收了 {collected} 个对象")

                if not _check_memory_before_operation(DOCUMENT_MAX_MEMORY_MB * 1.5):
                    print(f"[WARNING] 内存持续偏高，但继续处理...")

        print(f"[DOCX-CHUNKED] 所有块处理完成 ({processed_chunks}/{total_chunks})")

        doc.save(target_path)
        print(f"[DOCX-CHUNKED] 文档已保存到: {target_path}")

        del doc
        gc.collect()

        if ext == ".doc":
            return _convert_docx_to_doc(target_path, file_path)

        return _make_result({"message": f"Word文档创建成功: {file_path}", "method": "python-docx"})
    except MemoryError as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[ERROR] DOCX写入内存不足:\n{error_detail}")
        return _make_result(error=f"内存不足，文档内容过大 ({len(content)}字符)。建议：1) 减少文档内容 2) 分多次写入 3) 联系管理员调整内存限制")
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[ERROR] DOCX写入失败详情:\n{error_detail}")
        return _make_result(error=f"Word文档写入失败: {str(e)}")


def _convert_docx_to_doc(docx_path: str, target_doc_path: str) -> dict:
    """将 DOCX 转换为 DOC 格式（使用 LibreOffice）"""
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "doc",
             "--outdir", os.path.dirname(os.path.abspath(target_doc_path)), docx_path],
            capture_output=True,
            timeout=30,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        )
        os.unlink(docx_path)
        if result.returncode == 0:
            return _make_result({"message": f"DOC文档创建成功: {target_doc_path}"})
        else:
            return _make_result(error=f"DOC转换失败: {result.stderr}")
    except FileNotFoundError:
        os.unlink(docx_path)
        return _make_result(error="创建DOC需要LibreOffice支持，已生成DOCX版本")
    except subprocess.TimeoutExpired:
        os.unlink(docx_path)
        return _make_result(error="DOC转换超时")


def _docx_append(file_path: str, content: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return _make_result(error="缺少依赖: pip install python-docx")
    
    try:
        ext = _get_ext(file_path)
        actual_path = file_path
        
        if ext == ".doc":
            converted = _convert_doc_to_docx(file_path)
            if converted:
                actual_path = converted
            else:
                return _make_result(error="无法追加DOC格式，请转换为DOCX")
        
        doc = Document(actual_path)
        _markdown_to_docx_content(content, doc)
        doc.save(actual_path)
        
        if ext == ".doc":
            try:
                import subprocess
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "doc",
                     "--outdir", os.path.dirname(os.path.abspath(file_path)), actual_path],
                    capture_output=True, timeout=30
                )
                os.unlink(actual_path)
                if result.returncode == 0:
                    return _make_result({"message": "DOC追加成功"})
            except (FileNotFoundError, Exception):
                return _make_result(error="DOC转换失败")
        
        return _make_result({"message": "Word文档追加成功"})
    except Exception as e:
        return _make_result(error=f"Word文档追加失败: {str(e)}")


def _docx_update(file_path: str, target: str, content: str, field: Optional[str] = None) -> dict:
    try:
        from docx import Document
    except ImportError:
        return _make_result(error="缺少依赖: pip install python-docx")
    
    try:
        ext = _get_ext(file_path)
        actual_path = file_path
        cleanup = False
        
        if ext == ".doc":
            converted = _convert_doc_to_docx(file_path)
            if not converted:
                return _make_result(error="无法修改DOC格式")
            actual_path = converted
            cleanup = True
        
        doc = Document(actual_path)
        
        if field == "paragraph":
            try:
                para_idx = int(target)
                if 0 <= para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = content
                    else:
                        para.add_run(content)
                    doc.save(actual_path)
                    
                    if cleanup:
                        _save_as_doc(actual_path, file_path)
                        os.unlink(actual_path)
                    
                    return _make_result({"message": f"段落{para_idx}更新成功"})
                else:
                    return _make_result(error=f"段落索引超出范围(0-{len(doc.paragraphs)-1})")
            except ValueError:
                return _make_result(error="target必须是段落索引数字")
        
        elif field == "metadata":
            meta_fields = json.loads(content) if isinstance(content, str) else content
            props = doc.core_properties
            for k, v in meta_fields.items():
                if hasattr(props, k):
                    setattr(props, k, v)
            doc.save(actual_path)
            
            if cleanup:
                _save_as_doc(actual_path, file_path)
                os.unlink(actual_path)
            
            return _make_result({"message": "Word元数据更新成功"})
        
        if cleanup and os.path.exists(actual_path):
            os.unlink(actual_path)
        return _make_result(error=f"不支持的操作: field={field}, 支持 paragraph/metadata")
    except Exception as e:
        return _make_result(error=f"Word文档更新失败: {str(e)}")


def _save_as_doc(docx_path: str, output_doc: str) -> bool:
    try:
        import subprocess
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "doc",
             "--outdir", os.path.dirname(os.path.abspath(output_doc)), docx_path],
            capture_output=True, timeout=30
        )
        return True
    except Exception:
        return False


# ============================================================
# Excel XLS/XLSX 操作 (r/w/a/u)
# ============================================================

def _excel_read(file_path: str, start_idx: int = 0, max_length: int = 100000,
                include_metadata: bool = True) -> dict:
    try:
        import pandas as pd
    except ImportError:
        return _make_result(error="缺少依赖: pip install pandas")
    
    try:
        ext = _get_ext(file_path)
        
        if ext == ".xls":
            try:
                import xlrd
            except ImportError:
                return _make_result(error="缺少依赖: pip install xlrd")
        
        xl_file = pd.ExcelFile(file_path)
        sheet_names = xl_file.sheet_names
        
        all_content = []
        sheet_info = []
        
        for sheet_name in sheet_names:
            df = pd.read_excel(xl_file, sheet_name=sheet_name, header=None)
            
            rows_data = []
            for _, row in df.iterrows():
                row_values = [str(cell) if pd.notna(cell) else "" for cell in row]
                if any(v.strip() for v in row_values):
                    rows_data.append(" | ".join(row_values))
            
            sheet_content = f"## Sheet: {sheet_name}\n" + "\n".join(rows_data)
            all_content.append(sheet_content)
            
            sheet_info.append({
                "name": sheet_name,
                "rows": len(df),
                "cols": len(df.columns),
            })
        
        full_text = "\n\n".join(all_content)
        total_length = len(full_text)

        # 防御性处理：当 start_idx >= total_length 时，说明文档已经读完
        # ⚠️ 这个检查必须在切片之前！
        if start_idx >= total_length and total_length > 0:
            return _make_result({
                "content": "",
                "metadata": {},
                "structure": sheet_info,
                "total_length": total_length,
                "read_range": f"{total_length}-{total_length}",
                "truncated": False,
                "status": "end_of_file",
                "message": f"文档已读取完毕（总长度: {total_length}）"
            })

        end_idx = min(start_idx + max_length, total_length)
        content = full_text[start_idx:end_idx]

        metadata = {}
        if include_metadata:
            metadata = {
                "file_type": ext.lstrip("."),
                "sheet_count": len(sheet_names),
                "sheet_names": sheet_names,
            }
        
        return _make_result({
            "content": content,
            "metadata": metadata,
            "structure": sheet_info,
            "total_length": total_length,
            "read_range": f"{start_idx}-{end_idx}",
            "truncated": end_idx < total_length
        })
    except Exception as e:
        return _make_result(error=f"Excel读取失败: {str(e)}")


def _excel_write(file_path: str, data: Dict[str, list], metadata: Optional[dict] = None) -> dict:
    ext = _get_ext(file_path)
    
    if ext == ".xlsx":
        try:
            from openpyxl import Workbook
        except ImportError:
            return _make_result(error="缺少依赖: pip install openpyxl")
        
        try:
            wb = Workbook()
            default_sheet = wb.active
            
            for idx, (sheet_name, rows) in enumerate(data.items()):
                if idx == 0:
                    default_sheet.title = sheet_name
                    ws = default_sheet
                else:
                    ws = wb.create_sheet(title=sheet_name)
                
                for row_data in rows:
                    ws.append(row_data)
            
            wb.save(file_path)
            return _make_result({"message": f"XLSX创建成功: {file_path}", "sheets": list(data.keys())})
        except Exception as e:
            return _make_result(error=f"XLSX写入失败: {str(e)}")
    
    elif ext == ".xls":
        try:
            import xlwt
        except ImportError:
            return _make_result(error="缺少依赖: pip install xlwt")
        
        try:
            wb = xlwt.Workbook()
            
            for idx, (sheet_name, rows) in enumerate(data.items()):
                ws = wb.add_sheet(sheet_name[:31])
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, value in enumerate(row_data):
                        ws.write(row_idx, col_idx, value)
            
            wb.save(file_path)
            return _make_result({"message": f"XLS创建成功: {file_path}", "sheets": list(data.keys())})
        except Exception as e:
            return _make_result(error=f"XLS写入失败: {str(e)}")
    
    return _make_result(error=f"不支持的Excel格式: {ext}")


def _excel_append(file_path: str, data: Dict[str, list], target_sheet: Optional[str] = None) -> dict:
    ext = _get_ext(file_path)
    
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError:
            return _make_result(error="缺少依赖: pip install openpyxl")
        
        try:
            wb = load_workbook(file_path)
            
            for sheet_name, rows in data.items():
                actual_sheet = sheet_name
                if target_sheet:
                    actual_sheet = target_sheet
                
                if actual_sheet in wb.sheetnames:
                    ws = wb[actual_sheet]
                    start_row = ws.max_row + 1
                    
                    for row_data in rows:
                        for col_idx, value in enumerate(row_data, 1):
                            ws.cell(row=start_row, column=col_idx, value=value)
                        start_row += 1
                else:
                    ws = wb.create_sheet(title=actual_sheet)
                    for row_data in rows:
                        ws.append(row_data)
            
            wb.save(file_path)
            return _make_result({"message": "Excel数据追加成功"})
        except Exception as e:
            return _make_result(error=f"Excel追加失败: {str(e)}")
    
    elif ext == ".xls":
        try:
            import xlutils
            import xlrd
            import xlwt
        except ImportError:
            return _make_result(error="缺少依赖: pip install xlutils xlrd xlwt")
        
        try:
            rb = xlrd.open_workbook(file_path, formatting_info=True)
            wb = xlutils.copy.copy(rb)
            
            for sheet_name, rows in data.items():
                actual_sheet = sheet_name
                if target_sheet:
                    actual_sheet = target_sheet
                
                if actual_sheet in rb.sheet_names():
                    ws = wb.get_sheet(rb.sheet_names().index(actual_sheet))
                    start_row = ws.last_used_row + 1 if hasattr(ws, 'last_used_row') else rb.sheet_by_name(actual_sheet).nrows
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, value in enumerate(row_data):
                            ws.write(start_row + row_idx, col_idx, value)
                else:
                    ws = wb.add_sheet(actual_sheet)
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, value in enumerate(row_data):
                            ws.write(row_idx, col_idx, value)
            
            wb.save(file_path)
            return _make_result({"message": "Excel数据追加成功"})
        except Exception as e:
            return _make_result(error=f"Excel追加失败: {str(e)}")
    
    return _make_result(error=f"不支持的格式: {ext}")


def _excel_update(file_path: str, target: str, content: Any,
                  sheet_name: Optional[str] = None) -> dict:
    ext = _get_ext(file_path)
    
    try:
        cell_match = target.replace(" ", "").upper()
        import re
        match = re.match(r'([A-Z]+)(\d+)', cell_match)
        if not match:
            return _make_result(error="target格式错误，应为单元格坐标如'A1', 'B3'")
        
        col_str = match.group(1)
        row_num = int(match.group(2))
        
        col_num = 0
        for ch in col_str:
            col_num = col_num * 26 + (ord(ch) - ord('A') + 1)
    except Exception:
        return _make_result(error="target解析失败")
    
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError:
            return _make_result(error="缺少依赖: pip install openpyxl")
        
        try:
            wb = load_workbook(file_path)
            ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
            ws.cell(row=row_num, column=col_num, value=content)
            wb.save(file_path)
            return _make_result({"message": f"单元格{target}更新为: {content}"})
        except Exception as e:
            return _make_result(error=f"XLSX更新失败: {str(e)}")
    
    elif ext == ".xls":
        try:
            import xlutils
            import xlrd
        except ImportError:
            return _make_result(error="缺少依赖: pip install xlutils xlrd")
        
        try:
            rb = xlrd.open_workbook(file_path, formatting_info=True)
            wb = xlutils.copy.copy(rb)
            
            s_idx = rb.sheet_names().index(sheet_name) if sheet_name and sheet_name in rb.sheet_names() else 0
            ws = wb.get_sheet(s_idx)
            ws.write(row_num - 1, col_num - 1, content)
            wb.save(file_path)
            return _make_result({"message": f"单元格{target}更新为: {content}"})
        except Exception as e:
            return _make_result(error=f"XLS更新失败: {str(e)}")
    
    return _make_result(error=f"不支持的格式: {ext}")


# ============================================================
# 统一入口：execute_document (类似 fopen)
# ============================================================

def _redirect_to_workspace(file_path: str) -> str:
    if os.path.isabs(file_path):
        return file_path
    try:
        settings = get_settings_service()
        workspace_id = settings.get("current:workspace_id")
        session_id = settings.get("current:session_id")
        if workspace_id and session_id:
            workspace_dir = os.path.join("workspaces", session_id, workspace_id)
            os.makedirs(workspace_dir, exist_ok=True)
            redirected = os.path.join(workspace_dir, file_path)
            print(f"[PATH-REDIRECT] {file_path} -> {redirected}")
            return redirected
    except Exception as e:
        print(f"[PATH-REDIRECT] Warning: cannot get workspace info: {e}")
    return file_path


def execute_document(tool_args: dict, conversation_id: Optional[str] = None) -> dict:
    operation = tool_args.get("operation")
    file_path = tool_args.get("file_path")
    
    if not operation:
        return _make_result(error="缺少 operation 参数 (r|w|a|u|s)")
    
    if not file_path:
        return _make_result(error="缺少 file_path 参数")
    
    valid_ops = {"r", "w", "a", "u", "s"}
    if operation not in valid_ops:
        return _make_result(error=f"无效操作类型: {operation}，支持: {'|'.join(sorted(valid_ops))}")
    
    print(f"[Tool] document [{operation}] {file_path}")
    
    if operation != "w":
        if not os.path.exists(file_path):
            return _make_result(error=f"文件不存在: {file_path}")
        if not os.path.isfile(file_path):
            return _make_result(error=f"路径不是文件: {file_path}")
    
    ext = _get_ext(file_path)
    supported_formats = {".pdf", ".doc", ".docx", ".xls", ".xlsx"}
    
    if ext not in supported_formats:
        return _make_result(error=f"不支持的格式: {ext}，支持: {', '.join(sorted(supported_formats))}")
    
    settings = get_settings_service()
    try:
        use_llm_parsing = settings.get("agent_tools:pdf:use_llm_parsing")
    except KeyError:
        use_llm_parsing = True
    
    # ---- READ ----
    if operation == "r":
        # 防御性类型转换，防止 LLM 传入字符串类型的参数
        try:
            start_idx = int(tool_args.get("start_idx", 0))
        except (ValueError, TypeError):
            start_idx = 0
        try:
            max_length = int(tool_args.get("max_length", 100000))
        except (ValueError, TypeError):
            max_length = 100000
        include_metadata = tool_args.get("include_metadata", True)

        line_pagination_requested = "offset" in tool_args or "limit" in tool_args
        offset = None
        limit = None
        if line_pagination_requested:
            if ext not in {".doc", ".docx"}:
                return _make_result(
                    error="offset and limit are only supported for DOC/DOCX reads"
                )
            try:
                offset = int(tool_args.get("offset", 1))
                limit = int(tool_args.get("limit", 2000))
            except (ValueError, TypeError):
                return _make_result(error="offset and limit must be integers")
            if offset < 1:
                return _make_result(error="offset must be at least 1")
            if limit < 1 or limit > 2000:
                return _make_result(error="limit must be between 1 and 2000")

        if ext == ".pdf":
            result = _pdf_read(file_path, start_idx, max_length, include_metadata, use_llm_parsing)
        elif ext in {".doc", ".docx"}:
            result = _docx_read(
                file_path,
                start_idx,
                max_length,
                include_metadata,
                conversation_id,
                offset=offset,
                limit=limit,
            )
        elif ext in {".xls", ".xlsx"}:
            result = _excel_read(file_path, start_idx, max_length, include_metadata)
        else:
            result = _make_result(error=f"读取暂不支持: {ext}")
    
    # ---- WRITE ----
    elif operation == "w":
        content = tool_args.get("content", "")
        data = tool_args.get("data")
        metadata = tool_args.get("metadata")
        
        if ext == ".pdf":
            if not content:
                return _make_result(error="PDF写入需要content参数")
            result = _pdf_write(file_path, content, metadata)
        elif ext in {".doc", ".docx"}:
            if not content:
                return _make_result(error="Word写入需要content参数(Markdown文本)")
            result = _docx_write(file_path, content, metadata)
        elif ext in {".xls", ".xlsx"}:
            if not data:
                return _make_result(error="Excel写入需要data参数(JSON数组)")
            result = _excel_write(file_path, data, metadata)
        else:
            result = _make_result(error=f"写入暂不支持: {ext}")
    
    # ---- APPEND ----
    elif operation == "a":
        content = tool_args.get("content", "")
        data = tool_args.get("data")
        position = tool_args.get("position", "end")
        
        if ext == ".pdf":
            if not content:
                return _make_result(error="PDF追加需要content参数")
            result = _pdf_append(file_path, content)
        elif ext == ".docx":
            if not content:
                return _make_result(error="DOCX追加需要content参数")
            result = _docx_append(file_path, content)
        elif ext == ".doc":
            result = _make_result(error="DOC格式建议使用write模式覆盖或转为DOCX")
        elif ext in {".xls", ".xlsx"}:
            if not data:
                return _make_result(error="Excel追加需要data参数")
            target_sheet = tool_args.get("target") or tool_args.get("sheet_name")
            result = _excel_append(file_path, data, target_sheet)
        else:
            result = _make_result(error=f"追加暂不支持: {ext}")
    
    # ---- UPDATE ----
    elif operation == "u":
        target = tool_args.get("target")
        content = tool_args.get("content")
        field = tool_args.get("field")

        if not target:
            return _make_result(error="update操作需要target参数(定位信息)")

        if ext == ".pdf":
            result = _pdf_update(file_path, target, content, field or "metadata")
        elif ext in {".doc", ".docx"}:
            result = _docx_update(file_path, target, content, field or "paragraph")
        elif ext in {".xls", ".xlsx"}:
            if not content:
                return _make_result(error="Excel update需要content参数(新值)")
            sheet_name = tool_args.get("sheet_name")
            result = _excel_update(file_path, target, content, sheet_name)
        else:
            result = _make_result(error=f"更新暂不支持: {ext}")

    # ---- SEARCH (grep) ----
    elif operation == "s":
        pattern = tool_args.get("pattern")
        if not pattern:
            return _make_result(error="search操作需要pattern参数(搜索正则表达式)")

        case_sensitive = tool_args.get("case_sensitive", False)
        context = tool_args.get("context", 2)
        max_results = tool_args.get("max_results", 50)
        start_idx = tool_args.get("start_idx", 0)

        if ext == ".pdf":
            result = _pdf_search(
                file_path,
                pattern,
                case_sensitive,
                context,
                max_results,
                use_llm_parsing,
                start_idx,
            )
        elif ext in {".doc", ".docx"}:
            result = _docx_search(
                file_path,
                pattern,
                case_sensitive,
                context,
                max_results,
                conversation_id,
                start_idx,
            )
        elif ext in {".xls", ".xlsx"}:
            result = _excel_search(
                file_path,
                pattern,
                case_sensitive,
                context,
                max_results,
                start_idx,
            )
        else:
            result = _make_result(error=f"搜索暂不支持: {ext}")
    
    else:
        result = _make_result(error=f"未知操作: {operation}")
    
    if result.get("error") is None:
        print(f"[Tool] document [{operation}] 成功: {file_path}")
    else:
        print(f"[Tool] document [{operation}] 失败: {result['error']}")
    
    return result


DOCUMENT_TOOLS = {
    "document": ToolDefinition(
        name="document",
        description="统一文档操作，支持 PDF/DOC/DOCX/XLS/XLSX；operation: r=读 w=写 a=追加 u=修改 s=搜索(grep)。s 按命中行返回，行内命中汇总在 occurrences；返回上下文片段、字符偏移、段号、read_hint 和 next_start_idx",
        params='document:(仅支持doc/docx/pdf/xls/xlsx；写md/txt/json等文本文件请用write_file){"operation":"r|w|a|u|s(必填)","file_path":"(必填)","content":"(文本)","data":"(Excel)","pattern":"(搜索正则)","case_sensitive":false,"context":2,"max_results":50,"start_idx":"(r读取起点；s搜索起点，默认0；后续页传上次next_start_idx)"}',
        category="document",
        executor=execute_document
    )
}


def register_document_tools():
    for tool_name, tool_def in DOCUMENT_TOOLS.items():
        ToolRegistry.register(tool_def)
